from typing import Dict, List
from io import BytesIO
from datetime import datetime
from reportlab.lib.pagesizes import letter
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.lib.units import inch
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, PageBreak, Table, TableStyle
from reportlab.lib import colors
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from app.core.logging import logger

class ExportService:
    
    def export_to_pdf(self, brd_content: Dict, project_name: str) -> BytesIO:
        """Export BRD to PDF"""
        buffer = BytesIO()
        doc = SimpleDocTemplate(buffer, pagesize=letter)
        story = []
        styles = getSampleStyleSheet()
        
        # Title
        title_style = ParagraphStyle(
            'CustomTitle',
            parent=styles['Heading1'],
            fontSize=24,
            textColor=colors.HexColor('#1e40af'),
            spaceAfter=30,
            alignment=1  # Center
        )
        story.append(Paragraph(f"Business Requirements Document", title_style))
        story.append(Paragraph(f"{project_name}", title_style))
        story.append(Spacer(1, 0.5*inch))
        
        # Metadata
        meta_data = [
            ['Generated:', datetime.utcnow().strftime('%Y-%m-%d %H:%M UTC')],
            ['Version:', str(brd_content.get('version', '1.0'))],
        ]
        meta_table = Table(meta_data, colWidths=[2*inch, 4*inch])
        meta_table.setStyle(TableStyle([
            ('FONTNAME', (0, 0), (-1, -1), 'Helvetica'),
            ('FONTSIZE', (0, 0), (-1, -1), 10),
            ('TEXTCOLOR', (0, 0), (0, -1), colors.grey),
        ]))
        story.append(meta_table)
        story.append(Spacer(1, 0.5*inch))
        
        # Table of Contents
        story.append(Paragraph("Table of Contents", styles['Heading1']))
        toc_items = [
            "1. Executive Summary",
            "2. Business Objectives",
            "3. Functional Requirements",
            "4. Non-Functional Requirements",
            "5. Assumptions",
            "6. Success Metrics",
            "7. Timeline",
            "8. Risks"
        ]
        for item in toc_items:
            story.append(Paragraph(item, styles['Normal']))
        story.append(PageBreak())
        
        # Sections
        sections = [
            ('executive_summary', 'Executive Summary'),
            ('business_objectives', 'Business Objectives'),
            ('functional_requirements', 'Functional Requirements'),
            ('non_functional_requirements', 'Non-Functional Requirements'),
            ('assumptions', 'Assumptions'),
            ('success_metrics', 'Success Metrics'),
            ('timeline', 'Timeline'),
            ('risks', 'Risks')
        ]
        
        for section_key, section_title in sections:
            story.append(Paragraph(f"{sections.index((section_key, section_title)) + 1}. {section_title}", styles['Heading1']))
            story.append(Spacer(1, 0.2*inch))
            
            content = brd_content.get(section_key, [])
            if isinstance(content, list):
                for idx, item in enumerate(content, 1):
                    if isinstance(item, dict):
                        text = item.get('content', str(item))
                    else:
                        text = str(item)
                    story.append(Paragraph(f"{idx}. {text}", styles['Normal']))
                    story.append(Spacer(1, 0.1*inch))
            else:
                story.append(Paragraph(str(content), styles['Normal']))
            
            story.append(Spacer(1, 0.3*inch))
        
        doc.build(story)
        buffer.seek(0)
        return buffer
    
    def export_to_docx(self, brd_content: Dict, project_name: str) -> BytesIO:
        """Export BRD to DOCX"""
        doc = Document()
        
        # Title
        title = doc.add_heading('Business Requirements Document', 0)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        subtitle = doc.add_heading(project_name, level=1)
        subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Metadata
        doc.add_paragraph(f"Generated: {datetime.utcnow().strftime('%Y-%m-%d %H:%M UTC')}")
        doc.add_paragraph(f"Version: {brd_content.get('version', '1.0')}")
        doc.add_page_break()
        
        # Table of Contents
        doc.add_heading('Table of Contents', level=1)
        toc_items = [
            "1. Executive Summary",
            "2. Business Objectives",
            "3. Functional Requirements",
            "4. Non-Functional Requirements",
            "5. Assumptions",
            "6. Success Metrics",
            "7. Timeline",
            "8. Risks"
        ]
        for item in toc_items:
            doc.add_paragraph(item, style='List Number')
        doc.add_page_break()
        
        # Sections
        sections = [
            ('executive_summary', 'Executive Summary'),
            ('business_objectives', 'Business Objectives'),
            ('functional_requirements', 'Functional Requirements'),
            ('non_functional_requirements', 'Non-Functional Requirements'),
            ('assumptions', 'Assumptions'),
            ('success_metrics', 'Success Metrics'),
            ('timeline', 'Timeline'),
            ('risks', 'Risks')
        ]
        
        for section_key, section_title in sections:
            doc.add_heading(f"{sections.index((section_key, section_title)) + 1}. {section_title}", level=1)
            
            content = brd_content.get(section_key, [])
            if isinstance(content, list):
                for item in content:
                    if isinstance(item, dict):
                        text = item.get('content', str(item))
                    else:
                        text = str(item)
                    doc.add_paragraph(text, style='List Bullet')
            else:
                doc.add_paragraph(str(content))
        
        buffer = BytesIO()
        doc.save(buffer)
        buffer.seek(0)
        return buffer
    
    def export_to_markdown(self, brd_content: Dict, project_name: str) -> str:
        """Export BRD to Markdown"""
        md = []
        
        # Title
        md.append(f"# Business Requirements Document")
        md.append(f"## {project_name}")
        md.append("")
        md.append(f"**Generated:** {datetime.utcnow().strftime('%Y-%m-%d %H:%M UTC')}")
        md.append(f"**Version:** {brd_content.get('version', '1.0')}")
        md.append("")
        md.append("---")
        md.append("")
        
        # Table of Contents
        md.append("## Table of Contents")
        md.append("")
        md.append("1. [Executive Summary](#executive-summary)")
        md.append("2. [Business Objectives](#business-objectives)")
        md.append("3. [Functional Requirements](#functional-requirements)")
        md.append("4. [Non-Functional Requirements](#non-functional-requirements)")
        md.append("5. [Assumptions](#assumptions)")
        md.append("6. [Success Metrics](#success-metrics)")
        md.append("7. [Timeline](#timeline)")
        md.append("8. [Risks](#risks)")
        md.append("")
        md.append("---")
        md.append("")
        
        # Sections
        sections = [
            ('executive_summary', 'Executive Summary'),
            ('business_objectives', 'Business Objectives'),
            ('functional_requirements', 'Functional Requirements'),
            ('non_functional_requirements', 'Non-Functional Requirements'),
            ('assumptions', 'Assumptions'),
            ('success_metrics', 'Success Metrics'),
            ('timeline', 'Timeline'),
            ('risks', 'Risks')
        ]
        
        for section_key, section_title in sections:
            md.append(f"## {section_title}")
            md.append("")
            
            content = brd_content.get(section_key, [])
            if isinstance(content, list):
                for item in content:
                    if isinstance(item, dict):
                        text = item.get('content', str(item))
                    else:
                        text = str(item)
                    md.append(f"- {text}")
            else:
                md.append(str(content))
            
            md.append("")
        
        return "\n".join(md)

export_service = ExportService()
