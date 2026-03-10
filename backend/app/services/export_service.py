from typing import Dict
from io import BytesIO
from datetime import datetime
from reportlab.lib.pagesizes import letter
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.lib.units import inch
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, PageBreak, Table, TableStyle
from reportlab.lib import colors
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from app.core.logging import logger

class ExportService:
    
    def generate_pdf(self, brd_content: Dict, project_name: str) -> BytesIO:
        """Generate PDF from BRD content."""
        buffer = BytesIO()
        doc = SimpleDocTemplate(buffer, pagesize=letter,
                                rightMargin=72, leftMargin=72,
                                topMargin=72, bottomMargin=18)
        
        story = []
        styles = getSampleStyleSheet()
        
        # Custom styles
        title_style = ParagraphStyle(
            'CustomTitle',
            parent=styles['Heading1'],
            fontSize=24,
            textColor=colors.HexColor('#1e40af'),
            spaceAfter=30,
            alignment=1  # Center
        )
        
        heading_style = ParagraphStyle(
            'CustomHeading',
            parent=styles['Heading2'],
            fontSize=16,
            textColor=colors.HexColor('#1e40af'),
            spaceAfter=12,
            spaceBefore=12
        )
        
        # Title
        story.append(Paragraph(f"Business Requirements Document", title_style))
        story.append(Paragraph(f"{project_name}", styles['Heading2']))
        story.append(Spacer(1, 0.2*inch))
        
        # Metadata
        story.append(Paragraph(f"Generated: {datetime.now().strftime('%Y-%m-%d %H:%M')}", styles['Normal']))
        story.append(Spacer(1, 0.3*inch))
        
        # Table of Contents
        story.append(Paragraph("Table of Contents", heading_style))
        toc_data = [
            ["1.", "Executive Summary"],
            ["2.", "Business Objectives"],
            ["3.", "Functional Requirements"],
            ["4.", "Non-Functional Requirements"],
            ["5.", "Assumptions"],
            ["6.", "Risks"],
            ["7.", "Timeline"],
            ["8.", "Success Metrics"]
        ]
        toc_table = Table(toc_data, colWidths=[0.5*inch, 5*inch])
        toc_table.setStyle(TableStyle([
            ('FONTNAME', (0, 0), (-1, -1), 'Helvetica'),
            ('FONTSIZE', (0, 0), (-1, -1), 11),
            ('TEXTCOLOR', (0, 0), (-1, -1), colors.black),
            ('ALIGN', (0, 0), (-1, -1), 'LEFT'),
            ('BOTTOMPADDING', (0, 0), (-1, -1), 6),
        ]))
        story.append(toc_table)
        story.append(PageBreak())
        
        # Sections
        sections = [
            ("Executive Summary", brd_content.get('executive_summary', '')),
            ("Business Objectives", brd_content.get('business_objectives', [])),
            ("Functional Requirements", brd_content.get('functional_requirements', [])),
            ("Non-Functional Requirements", brd_content.get('non_functional_requirements', [])),
            ("Assumptions", brd_content.get('assumptions', [])),
            ("Risks", brd_content.get('risks', [])),
            ("Timeline", brd_content.get('timeline', [])),
            ("Success Metrics", brd_content.get('success_metrics', []))
        ]
        
        for section_title, section_content in sections:
            story.append(Paragraph(section_title, heading_style))
            
            if isinstance(section_content, str):
                story.append(Paragraph(section_content, styles['Normal']))
            elif isinstance(section_content, list):
                for idx, item in enumerate(section_content, 1):
                    if isinstance(item, dict):
                        text = f"{idx}. {item.get('content', '')}"
                    else:
                        text = f"{idx}. {item}"
                    story.append(Paragraph(text, styles['Normal']))
                    story.append(Spacer(1, 0.1*inch))
            
            story.append(Spacer(1, 0.2*inch))
        
        doc.build(story)
        buffer.seek(0)
        return buffer
    
    def generate_docx(self, brd_content: Dict, project_name: str) -> BytesIO:
        """Generate DOCX from BRD content."""
        doc = Document()
        
        # Title
        title = doc.add_heading('Business Requirements Document', 0)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        subtitle = doc.add_heading(project_name, level=1)
        subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Metadata
        p = doc.add_paragraph()
        p.add_run(f"Generated: {datetime.now().strftime('%Y-%m-%d %H:%M')}").italic = True
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        doc.add_page_break()
        
        # Table of Contents
        doc.add_heading('Table of Contents', level=1)
        toc_items = [
            "1. Executive Summary",
            "2. Business Objectives",
            "3. Functional Requirements",
            "4. Non-Functional Requirements",
            "5. Assumptions",
            "6. Risks",
            "7. Timeline",
            "8. Success Metrics"
        ]
        for item in toc_items:
            doc.add_paragraph(item, style='List Number')
        
        doc.add_page_break()
        
        # Sections
        sections = [
            ("Executive Summary", brd_content.get('executive_summary', '')),
            ("Business Objectives", brd_content.get('business_objectives', [])),
            ("Functional Requirements", brd_content.get('functional_requirements', [])),
            ("Non-Functional Requirements", brd_content.get('non_functional_requirements', [])),
            ("Assumptions", brd_content.get('assumptions', [])),
            ("Risks", brd_content.get('risks', [])),
            ("Timeline", brd_content.get('timeline', [])),
            ("Success Metrics", brd_content.get('success_metrics', []))
        ]
        
        for section_title, section_content in sections:
            doc.add_heading(section_title, level=1)
            
            if isinstance(section_content, str):
                doc.add_paragraph(section_content)
            elif isinstance(section_content, list):
                for idx, item in enumerate(section_content, 1):
                    if isinstance(item, dict):
                        text = item.get('content', '')
                        p = doc.add_paragraph(text, style='List Number')
                        if item.get('priority'):
                            p.add_run(f" [Priority: {item['priority']}]").bold = True
                    else:
                        doc.add_paragraph(str(item), style='List Number')
        
        buffer = BytesIO()
        doc.save(buffer)
        buffer.seek(0)
        return buffer
    
    def generate_markdown(self, brd_content: Dict, project_name: str) -> str:
        """Generate Markdown from BRD content."""
        md = f"# Business Requirements Document\n\n"
        md += f"## {project_name}\n\n"
        md += f"*Generated: {datetime.now().strftime('%Y-%m-%d %H:%M')}*\n\n"
        md += "---\n\n"
        
        # Table of Contents
        md += "## Table of Contents\n\n"
        md += "1. [Executive Summary](#executive-summary)\n"
        md += "2. [Business Objectives](#business-objectives)\n"
        md += "3. [Functional Requirements](#functional-requirements)\n"
        md += "4. [Non-Functional Requirements](#non-functional-requirements)\n"
        md += "5. [Assumptions](#assumptions)\n"
        md += "6. [Risks](#risks)\n"
        md += "7. [Timeline](#timeline)\n"
        md += "8. [Success Metrics](#success-metrics)\n\n"
        md += "---\n\n"
        
        # Sections
        sections = [
            ("Executive Summary", brd_content.get('executive_summary', '')),
            ("Business Objectives", brd_content.get('business_objectives', [])),
            ("Functional Requirements", brd_content.get('functional_requirements', [])),
            ("Non-Functional Requirements", brd_content.get('non_functional_requirements', [])),
            ("Assumptions", brd_content.get('assumptions', [])),
            ("Risks", brd_content.get('risks', [])),
            ("Timeline", brd_content.get('timeline', [])),
            ("Success Metrics", brd_content.get('success_metrics', []))
        ]
        
        for section_title, section_content in sections:
            anchor = section_title.lower().replace(' ', '-')
            md += f"## {section_title}\n\n"
            
            if isinstance(section_content, str):
                md += f"{section_content}\n\n"
            elif isinstance(section_content, list):
                for idx, item in enumerate(section_content, 1):
                    if isinstance(item, dict):
                        text = item.get('content', '')
                        priority = item.get('priority', '')
                        md += f"{idx}. {text}"
                        if priority:
                            md += f" **[Priority: {priority}]**"
                        md += "\n"
                    else:
                        md += f"{idx}. {item}\n"
                md += "\n"
        
        return md

export_service = ExportService()
