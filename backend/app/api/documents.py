from fastapi import APIRouter, Depends, HTTPException
from fastapi.responses import StreamingResponse
from sqlalchemy.orm import Session
from typing import List
from app.database import get_db
from app.models import Project, Document, ExtractedItem, ExtractedRequirement, Conflict
from app.schemas import DocumentResponse
from app.services.generator import DocumentGenerator
from app.auth import require_editor, require_viewer
import io
import json

router = APIRouter()
generator = DocumentGenerator()

@router.post("/{project_id}/generate", response_model=DocumentResponse)
def generate_document(
    project_id: int, 
    current_user: dict = Depends(require_editor()),  # Editor or Admin can generate
    db: Session = Depends(get_db)
):
    """
    Generate BRD document for a project.
    Requires: Editor or Admin role.
    """
    try:
        print(f"Generate BRD called with project_id: {project_id}")
        
        project = db.query(Project).filter(Project.id == project_id).first()
        if not project:
            raise HTTPException(status_code=404, detail="Project not found")
        
        # Check both ExtractedRequirement (new Smart Extract) and ExtractedItem (legacy)
        requirements = db.query(ExtractedRequirement).filter(
            ExtractedRequirement.project_id == project_id
        ).all()
        
        legacy_items = db.query(ExtractedItem).filter(
            ExtractedItem.project_id == project_id
        ).all()
        
        print(f"Requirements found: {len(requirements)}")
        print(f"Legacy items found: {len(legacy_items)}")
        
        # Use whichever has data
        items = requirements if requirements else legacy_items
        
        if not items:
            raise HTTPException(status_code=400, detail="No requirements found for this project. Please extract requirements first.")
        
        conflicts = db.query(Conflict).filter(Conflict.project_id == project_id).all()
        
        print(f"Generating BRD with {len(items)} requirements...")
        brd_content = generator.generate_brd(project.name, items, project.sources, conflicts)
        print(f"BRD content generated successfully")
        
        # Get latest version number
        latest_doc = db.query(Document).filter(
            Document.project_id == project_id
        ).order_by(Document.version.desc()).first()
        
        version = (latest_doc.version + 1) if latest_doc else 1
        
        document = Document(
            project_id=project_id,
            version=version,
            content=brd_content
        )
        db.add(document)
        db.commit()
        db.refresh(document)
        
        print(f"BRD generated successfully with {len(items)} requirements, version {version}")
        
        return document
    
    except HTTPException:
        # Re-raise HTTP exceptions
        raise
    except Exception as e:
        # Log and return detailed error
        print(f"BRD GENERATION ERROR: {str(e)}")
        import traceback
        traceback.print_exc()
        raise HTTPException(
            status_code=500, 
            detail=f"Failed to generate BRD: {str(e)}"
        )

@router.get("/{project_id}/documents", response_model=List[DocumentResponse])
def list_documents(
    project_id: int, 
    current_user: dict = Depends(require_viewer()),  # All users can list
    db: Session = Depends(get_db)
):
    """
    List all document versions for a project.
    Requires: Any authenticated user.
    """
    documents = db.query(Document).filter(
        Document.project_id == project_id
    ).order_by(Document.version.desc()).all()
    return documents

@router.get("/document/{document_id}", response_model=DocumentResponse)
def get_document(
    document_id: int, 
    current_user: dict = Depends(require_viewer()),  # All users can read
    db: Session = Depends(get_db)
):
    """
    Get a specific document version.
    Requires: Any authenticated user.
    """
    document = db.query(Document).filter(Document.id == document_id).first()
    if not document:
        raise HTTPException(status_code=404, detail="Document not found")
    return document

@router.get("/{project_id}/traceability")
def get_traceability_matrix(
    project_id: int, 
    current_user: dict = Depends(require_viewer()),  # All users can view
    db: Session = Depends(get_db)
):
    """
    Generate requirement traceability matrix.
    Requires: Any authenticated user.
    """
    # Check both ExtractedRequirement (new Smart Extract) and ExtractedItem (legacy)
    requirements = db.query(ExtractedRequirement).filter(
        ExtractedRequirement.project_id == project_id
    ).all()
    
    legacy_items = db.query(ExtractedItem).filter(
        ExtractedItem.project_id == project_id
    ).all()
    
    items = requirements if requirements else legacy_items
    
    if not items:
        raise HTTPException(status_code=404, detail="No requirements found")
    
    matrix = generator.generate_traceability_matrix(items)
    return {"matrix": matrix}

@router.get("/{project_id}/export/pdf")
def export_pdf(
    project_id: int, 
    current_user: dict = Depends(require_viewer()),
    db: Session = Depends(get_db)
):
    """
    Export BRD as professional PDF document.
    Requires: Any authenticated user.
    """
    from reportlab.lib.pagesizes import A4, letter
    from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
    from reportlab.lib.units import inch
    from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, PageBreak, Table, TableStyle
    from reportlab.lib import colors
    from reportlab.lib.enums import TA_CENTER, TA_LEFT
    
    # Get project and document
    project = db.query(Project).filter(Project.id == project_id).first()
    if not project:
        raise HTTPException(status_code=404, detail="Project not found")
    
    latest_doc = db.query(Document).filter(
        Document.project_id == project_id
    ).order_by(Document.version.desc()).first()
    
    if not latest_doc:
        raise HTTPException(status_code=404, detail="No document found. Please generate BRD first.")
    
    # Create PDF buffer
    buffer = io.BytesIO()
    doc = SimpleDocTemplate(buffer, pagesize=letter, topMargin=0.75*inch, bottomMargin=0.75*inch)
    
    # Container for PDF elements
    story = []
    styles = getSampleStyleSheet()
    
    # Custom styles
    title_style = ParagraphStyle(
        'CustomTitle',
        parent=styles['Heading1'],
        fontSize=24,
        textColor=colors.HexColor('#1e40af'),
        spaceAfter=30,
        alignment=TA_CENTER
    )
    
    heading_style = ParagraphStyle(
        'CustomHeading',
        parent=styles['Heading2'],
        fontSize=16,
        textColor=colors.HexColor('#1e40af'),
        spaceAfter=12,
        spaceBefore=12
    )
    
    # Title
    content = latest_doc.content
    story.append(Paragraph(content.get('title', 'Business Requirements Document'), title_style))
    story.append(Spacer(1, 0.2*inch))
    
    # Metadata
    story.append(Paragraph(f"<b>Project:</b> {project.name}", styles['Normal']))
    story.append(Paragraph(f"<b>Generated:</b> {content.get('generated_at', 'N/A')}", styles['Normal']))
    story.append(Paragraph(f"<b>Version:</b> {content.get('version', '1')}", styles['Normal']))
    story.append(Spacer(1, 0.3*inch))
    
    # Sections
    sections = content.get('sections', {})
    
    for section_name, section_data in sections.items():
        if section_data is None:
            continue
        
        title = section_name.replace('_', ' ').title()
        story.append(Paragraph(title, heading_style))
        
        if isinstance(section_data, dict):
            # Executive summary or text content
            if 'content' in section_data:
                story.append(Paragraph(section_data['content'], styles['Normal']))
                story.append(Spacer(1, 0.1*inch))
            
            # Statistics
            if 'statistics' in section_data:
                stats = section_data['statistics']
                stats_text = " | ".join([f"<b>{k.replace('_', ' ').title()}:</b> {v}" for k, v in stats.items()])
                story.append(Paragraph(stats_text, styles['Normal']))
                story.append(Spacer(1, 0.1*inch))
            
            # Items (requirements)
            if 'items' in section_data and section_data['items']:
                for item in section_data['items']:
                    req_text = f"<b>{item.get('id', 'N/A')}:</b> {item.get('content', '')}"
                    if 'priority' in item:
                        req_text += f" <i>(Priority: {item['priority']})</i>"
                    story.append(Paragraph(req_text, styles['Normal']))
                    story.append(Spacer(1, 0.05*inch))
        
        story.append(Spacer(1, 0.2*inch))
    
    # Build PDF
    doc.build(story)
    buffer.seek(0)
    
    return StreamingResponse(
        buffer,
        media_type="application/pdf",
        headers={"Content-Disposition": f"attachment; filename=BRD_{project.name.replace(' ', '_')}_v{latest_doc.version}.pdf"}
    )

@router.get("/{project_id}/export/word")
def export_word(
    project_id: int, 
    current_user: dict = Depends(require_viewer()),
    db: Session = Depends(get_db)
):
    """
    Export BRD as editable Word document (.docx).
    Requires: Any authenticated user.
    """
    from docx import Document as DocxDocument
    from docx.shared import Inches, Pt, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    
    # Get project and document
    project = db.query(Project).filter(Project.id == project_id).first()
    if not project:
        raise HTTPException(status_code=404, detail="Project not found")
    
    latest_doc = db.query(Document).filter(
        Document.project_id == project_id
    ).order_by(Document.version.desc()).first()
    
    if not latest_doc:
        raise HTTPException(status_code=404, detail="No document found. Please generate BRD first.")
    
    # Create Word document
    doc = DocxDocument()
    content = latest_doc.content
    
    # Title
    title = doc.add_heading(content.get('title', 'Business Requirements Document'), level=0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # Metadata
    doc.add_paragraph(f"Project: {project.name}")
    doc.add_paragraph(f"Generated: {content.get('generated_at', 'N/A')}")
    doc.add_paragraph(f"Version: {content.get('version', '1')}")
    doc.add_paragraph()
    
    # Sections
    sections = content.get('sections', {})
    
    for section_name, section_data in sections.items():
        if section_data is None:
            continue
        
        title = section_name.replace('_', ' ').title()
        doc.add_heading(title, level=1)
        
        if isinstance(section_data, dict):
            # Text content
            if 'content' in section_data:
                doc.add_paragraph(section_data['content'])
            
            # Statistics
            if 'statistics' in section_data:
                stats = section_data['statistics']
                stats_para = doc.add_paragraph()
                for k, v in stats.items():
                    stats_para.add_run(f"{k.replace('_', ' ').title()}: ").bold = True
                    stats_para.add_run(f"{v}  ")
            
            # Items (requirements)
            if 'items' in section_data and section_data['items']:
                for item in section_data['items']:
                    para = doc.add_paragraph(style='List Bullet')
                    para.add_run(f"{item.get('id', 'N/A')}: ").bold = True
                    para.add_run(item.get('content', ''))
                    if 'priority' in item:
                        para.add_run(f" (Priority: {item['priority']})").italic = True
        
        doc.add_paragraph()
    
    # Save to buffer
    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    
    return StreamingResponse(
        buffer,
        media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        headers={"Content-Disposition": f"attachment; filename=BRD_{project.name.replace(' ', '_')}_v{latest_doc.version}.docx"}
    )

@router.get("/{project_id}/export/excel")
def export_excel(
    project_id: int, 
    current_user: dict = Depends(require_viewer()),
    db: Session = Depends(get_db)
):
    """
    Export requirements as structured Excel spreadsheet (.xlsx).
    Requires: Any authenticated user.
    """
    from openpyxl import Workbook
    from openpyxl.styles import Font, PatternFill, Alignment
    
    # Get project and requirements
    project = db.query(Project).filter(Project.id == project_id).first()
    if not project:
        raise HTTPException(status_code=404, detail="Project not found")
    
    # Get requirements from both tables
    requirements = db.query(ExtractedRequirement).filter(
        ExtractedRequirement.project_id == project_id
    ).all()
    
    legacy_items = db.query(ExtractedItem).filter(
        ExtractedItem.project_id == project_id
    ).all()
    
    items = requirements if requirements else legacy_items
    
    if not items:
        raise HTTPException(status_code=404, detail="No requirements found")
    
    # Create workbook
    wb = Workbook()
    ws = wb.active
    ws.title = "Requirements"
    
    # Header styling
    header_fill = PatternFill(start_color="1e40af", end_color="1e40af", fill_type="solid")
    header_font = Font(bold=True, color="FFFFFF", size=12)
    
    # Headers
    headers = ["ID", "Category/Type", "Requirement", "Priority", "Confidence", "Created"]
    ws.append(headers)
    
    # Style headers
    for cell in ws[1]:
        cell.fill = header_fill
        cell.font = header_font
        cell.alignment = Alignment(horizontal="center", vertical="center")
    
    # Add requirements
    for idx, item in enumerate(items, start=1):
        # Determine type
        if isinstance(item, ExtractedRequirement):
            req_type = item.requirement_type
            priority = item.priority
            confidence = f"{item.confidence_score}%"
        else:
            req_type = item.item_type.value if hasattr(item.item_type, 'value') else str(item.item_type)
            priority = "High" if item.confidence_score > 0.8 else "Medium" if item.confidence_score > 0.5 else "Low"
            confidence = f"{int(item.confidence_score * 100)}%"
        
        ws.append([
            f"REQ-{item.id}",
            req_type.replace('_', ' ').title(),
            item.content,
            priority.capitalize(),
            confidence,
            item.created_at.strftime("%Y-%m-%d") if item.created_at else "N/A"
        ])
    
    # Adjust column widths
    ws.column_dimensions['A'].width = 12
    ws.column_dimensions['B'].width = 20
    ws.column_dimensions['C'].width = 60
    ws.column_dimensions['D'].width = 12
    ws.column_dimensions['E'].width = 12
    ws.column_dimensions['F'].width = 15
    
    # Save to buffer
    buffer = io.BytesIO()
    wb.save(buffer)
    buffer.seek(0)
    
    return StreamingResponse(
        buffer,
        media_type="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
        headers={"Content-Disposition": f"attachment; filename=Requirements_{project.name.replace(' ', '_')}.xlsx"}
    )


@router.get("/test-gemini")
async def test_gemini():
    """Test endpoint to verify Gemini API integration."""
    from google import genai
    from app.core.config import settings
    
    # Check if API key is configured
    if not settings.GEMINI_API_KEY or settings.GEMINI_API_KEY == "YOUR_GEMINI_API_KEY_HERE":
        return {
            "status": "error",
            "message": "GEMINI_API_KEY not configured in .env file"
        }
    
    try:
        # Initialize client
        client = genai.Client(api_key=settings.GEMINI_API_KEY)
        
        # Test API call
        response = client.models.generate_content(
            model="models/gemini-2.0-flash",
            contents="Say: Gemini integration successful"
        )
        
        return {
            "status": "success",
            "message": "Gemini API is working correctly",
            "api_key_length": len(settings.GEMINI_API_KEY),
            "model": "models/gemini-2.0-flash",
            "result": response.text
        }
    except Exception as e:
        return {
            "status": "error",
            "message": f"Gemini API error: {str(e)}",
            "error_type": type(e).__name__
        }
